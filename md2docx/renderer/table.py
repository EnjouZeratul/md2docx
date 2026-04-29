"""Table renderer with three-line table support."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from md2docx.parser import TableElement
from md2docx.formatter import set_run_font


class TableRenderer:
    """Renders tables with three-line style."""

    def __init__(self, border_top_weight: float = 1.5,
                 border_header_weight: float = 1.0,
                 border_bottom_weight: float = 1.5,
                 border_color: str = "000000"):
        self.border_top_weight = border_top_weight
        self.border_header_weight = border_header_weight
        self.border_bottom_weight = border_bottom_weight
        self.border_color = border_color

    def render(self, doc: Document, table_elem: TableElement) -> None:
        """Render table element to document."""
        if not table_elem.headers and not table_elem.rows:
            return

        if table_elem.caption:
            caption_para = doc.add_paragraph()
            caption_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption_para.add_run(table_elem.caption)
            set_run_font(run, font_name='宋体', size=9)

        num_cols = len(table_elem.headers) if table_elem.headers else len(table_elem.rows[0]) if table_elem.rows else 0
        if num_cols == 0:
            return

        num_rows = len(table_elem.rows) + (1 if table_elem.headers else 0)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        if table_elem.headers:
            header_row = table.rows[0]
            for i, header in enumerate(table_elem.headers):
                cell = header_row.cells[i]
                cell.text = header
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True

        start_row = 1 if table_elem.headers else 0
        for i, row_data in enumerate(table_elem.rows):
            row = table.rows[start_row + i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = cell_text
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._apply_three_line_style(table)

    def _apply_three_line_style(self, table) -> None:
        """Apply three-line table style."""
        def set_cell_border(cell, border_type: str, weight: float):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(int(weight * 8)))
            border.set(qn('w:color'), self.border_color)
            tcBorders.append(border)
            tcPr.append(tcBorders)

        if table.rows:
            first_row = table.rows[0]
            for cell in first_row.cells:
                set_cell_border(cell, 'top', self.border_top_weight)

            if len(table.rows) > 1:
                for cell in first_row.cells:
                    set_cell_border(cell, 'bottom', self.border_header_weight)

            last_row = table.rows[-1]
            for cell in last_row.cells:
                set_cell_border(cell, 'bottom', self.border_bottom_weight)

        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for side in ['left', 'right']:
                    border = OxmlElement(f'w:{side}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
