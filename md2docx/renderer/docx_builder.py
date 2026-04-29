"""DOCX document builder.

Coordinates rendering and assembles final document.
"""

from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from md2docx.config import Config
from md2docx.parser import (
    Element, HeadingElement, ParagraphElement,
    MermaidBlockElement, FormulaBlockElement,
    TableElement, CodeBlockElement, ImageElement,
    PageBreakElement, AbstractElement
)
from md2docx.formatter import (
    set_run_font, set_para_spacing, set_outline_level
)
from md2docx.renderer.mermaid import MermaidRenderer
from md2docx.renderer.formula import FormulaRenderer
from md2docx.renderer.table import TableRenderer


class DocxBuilder:
    """Builds DOCX document from parsed elements."""

    def __init__(self, config: Config, verbose: bool = False):
        self.config = config
        self.verbose = verbose

        self.mermaid_renderer = MermaidRenderer()
        self.formula_renderer = FormulaRenderer(
            chinese_replacements=config.formula.chinese_replacements
        )
        self.table_renderer = TableRenderer(
            border_top_weight=config.table.border_top_weight,
            border_header_weight=config.table.border_header_weight,
            border_bottom_weight=config.table.border_bottom_weight,
            border_color=config.table.border_color
        )

    def build(self, elements: List[Element], output_path: str) -> None:
        """Build DOCX from element list."""
        doc = Document()

        self._setup_page(doc)

        for elem in elements:
            self._add_element(doc, elem)

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

    def _setup_page(self, doc: Document) -> None:
        """Setup page properties."""
        section = doc.sections[0]

        if self.config.page.size == "A4":
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)

        section.left_margin = Cm(float(self.config.page.margins.left.replace('cm', '')))
        section.right_margin = Cm(float(self.config.page.margins.right.replace('cm', '')))
        section.top_margin = Cm(float(self.config.page.margins.top.replace('cm', '')))
        section.bottom_margin = Cm(float(self.config.page.margins.bottom.replace('cm', '')))

    def _add_element(self, doc: Document, elem: Element) -> None:
        """Add element to document."""
        if isinstance(elem, HeadingElement):
            self._add_heading(doc, elem)
        elif isinstance(elem, ParagraphElement):
            self._add_paragraph(doc, elem)
        elif isinstance(elem, MermaidBlockElement):
            self._add_mermaid(doc, elem)
        elif isinstance(elem, FormulaBlockElement):
            self._add_formula(doc, elem)
        elif isinstance(elem, TableElement):
            self._add_table(doc, elem)
        elif isinstance(elem, CodeBlockElement):
            self._add_code(doc, elem)
        elif isinstance(elem, ImageElement):
            self._add_image(doc, elem)
        elif isinstance(elem, PageBreakElement):
            doc.add_page_break()

    def _add_heading(self, doc: Document, elem: HeadingElement) -> None:
        """Add heading element."""
        para = doc.add_paragraph()

        if elem.level == 2:  # 章
            style = self.config.headings.chapter
            outline_level = 0
        elif elem.level == 3:  # 节
            style = self.config.headings.section
            outline_level = 1
        else:  # 小节
            style = self.config.headings.subsection
            outline_level = 2

        run = para.add_run(elem.text)
        set_run_font(run, font_name=self.config.fonts.heading,
                    size=style.size, bold=style.bold)

        align = WD_ALIGN_PARAGRAPH.CENTER if style.align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.alignment = align

        set_outline_level(para, outline_level)
        set_para_spacing(para, space_before=12, space_after=6)

    def _add_paragraph(self, doc: Document, elem: ParagraphElement) -> None:
        """Add paragraph element."""
        para = doc.add_paragraph()

        if elem.has_inline_math:
            import re
            parts = re.split(r'(\$[^$]+\$)', elem.text)
            for part in parts:
                if part.startswith('$') and part.endswith('$'):
                    run = para.add_run(part[1:-1])
                    run.font.italic = True
                    set_run_font(run, font_name=self.config.fonts.formula,
                                size=self.config.body.size, italic=True)
                else:
                    if part:
                        run = para.add_run(part)
                        set_run_font(run, font_name=self.config.fonts.default,
                                    size=self.config.body.size)
        else:
            run = para.add_run(elem.text)
            set_run_font(run, font_name=self.config.fonts.default,
                        size=self.config.body.size)

        indent_value = float(self.config.body.first_line_indent.replace('cm', ''))
        set_para_spacing(para, line_spacing=self.config.body.line_spacing,
                        indent_cm=indent_value)

    def _add_mermaid(self, doc: Document, elem: MermaidBlockElement) -> None:
        """Add mermaid diagram."""
        if elem.caption:
            caption_para = doc.add_paragraph()
            caption_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption_para.add_run(elem.caption)
            set_run_font(run, size=self.config.figures.caption_size)

        png_data = self.mermaid_renderer.render(elem.code)

        if png_data:
            import io
            image_stream = io.BytesIO(png_data)
            para = doc.add_paragraph()
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            max_width_cm = float(self.config.figures.max_width.replace('cm', ''))
            run.add_picture(image_stream, width=Cm(max_width_cm))
        else:
            para = doc.add_paragraph()
            run = para.add_run(f"【{elem.caption}: Mermaid 图表渲染失败】")
            run.font.italic = True

    def _add_formula(self, doc: Document, elem: FormulaBlockElement) -> None:
        """Add formula block."""
        png_data = self.formula_renderer.render(elem.latex)

        if png_data:
            import io
            image_stream = io.BytesIO(png_data)

            para = doc.add_paragraph()
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_stream)

            if elem.tag:
                tag_para = doc.add_paragraph()
                tag_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = tag_para.add_run(f"({elem.tag})")
                set_run_font(run, size=10)
        else:
            para = doc.add_paragraph()
            run = para.add_run(f"[公式: {elem.latex}]")
            run.font.italic = True

    def _add_table(self, doc: Document, elem: TableElement) -> None:
        """Add table element."""
        self.table_renderer.render(doc, elem)

    def _add_code(self, doc: Document, elem: CodeBlockElement) -> None:
        """Add code block."""
        para = doc.add_paragraph()

        code_text = elem.code
        if elem.language:
            code_text = f"[{elem.language}]\n{elem.code}"

        run = para.add_run(code_text)
        set_run_font(run, font_name=self.config.fonts.code, size=9)
        para.paragraph_format.left_indent = Cm(0.5)

    def _add_image(self, doc: Document, elem: ImageElement) -> None:
        """Add image element."""
        import os

        if not os.path.exists(elem.path):
            para = doc.add_paragraph()
            run = para.add_run(f"【图片不存在: {elem.path}】")
            run.font.italic = True
            return

        if elem.caption:
            caption_para = doc.add_paragraph()
            caption_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption_para.add_run(elem.caption)
            set_run_font(run, size=self.config.figures.caption_size)

        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            run = para.add_run()
            max_width_cm = float(self.config.figures.max_width.replace('cm', ''))
            run.add_picture(elem.path, width=Cm(max_width_cm))
        except Exception as e:
            run = para.add_run(f"【图片加载失败: {elem.path}】")
            run.font.italic = True
