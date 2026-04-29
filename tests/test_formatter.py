# tests/test_formatter.py
import pytest
from docx import Document
from md2docx.formatter import (
    set_run_font, set_para_spacing, set_outline_level,
    add_header_border
)


class TestFormatter:
    def test_set_run_font(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("测试文本")

        set_run_font(run, font_name='宋体', bold=True, size=12, italic=False)

        assert run.font.bold == True
        assert run.font.name == '宋体'

    def test_set_para_spacing(self):
        doc = Document()
        para = doc.add_paragraph("测试段落")

        set_para_spacing(para, space_before=6, space_after=6,
                        line_spacing=1.5, indent_cm=0.74, align='left')

        assert para.paragraph_format.line_spacing is not None

    def test_set_outline_level(self):
        doc = Document()
        para = doc.add_paragraph("标题")

        # Should not raise
        set_outline_level(para, level=0)

    def test_add_header_border(self):
        doc = Document()
        section = doc.sections[0]
        header = section.header
        para = header.paragraphs[0]
        para.text = "页眉文本"

        # Should not raise
        add_header_border(para)
